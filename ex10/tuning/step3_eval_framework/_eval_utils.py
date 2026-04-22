"""step3 — 평가 실행 보조 함수 (완성 코드).

벡터DB 구축, LLM 답변 생성, 테스트 질문 로드 등을 제공한다.
실제 검색·리랭크·쿼리변환은 pipelines.py의 부품이 strategies.py를 통해 주입된다.
"""

import json
import os
from pathlib import Path
from typing import Callable

import chromadb
from chromadb.config import Settings
from dotenv import load_dotenv

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent.parent.parent
DATA_DIR = BASE_DIR / "data"
CHROMA_DIR = DATA_DIR / "chroma_db"
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "qwen2.5:7b")


def _parse_docx_as_pages(path: Path) -> list[str]:
    """DOCX 전체를 한 '페이지'로 압축해 list[str]로 반환한다.

    ex07 _parse_and_chunk_docs와 동일한 방식: Heading/Title/List Bullet
    스타일을 markdown으로 치환하고 표는 markdown table로 변환한다.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name
        if style_name == "Title":
            parts.append(f"# {text}")
        elif style_name.startswith("Heading"):
            level_str = style_name.replace("Heading", "").strip()
            try:
                level = int(level_str)
            except ValueError:
                level = 2
            parts.append(f"{'#' * level} {text}")
        elif style_name == "List Bullet":
            parts.append(f"- {text}")
        else:
            parts.append(text)
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            parts.append("| " + " | ".join(row_data) + " |")
            if i == 0:
                parts.append("| " + " | ".join(["---"] * len(row_data)) + " |")
    full = "\n\n".join(parts)
    return [full] if full else []


def _parse_xlsx_as_pages(path: Path) -> list[str]:
    """XLSX를 시트별 1페이지로 쪼개 list[str]로 반환한다.

    ex07 로직과 동일: 각 시트를 [시트: 이름] 제목 + markdown table로 변환.
    """
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True)
    pages: list[str] = []
    for name in wb.sheetnames:
        ws = wb[name]
        rows: list[list[str]] = []
        for row in ws.iter_rows():
            cell_values = [
                str(c.value).strip()
                for c in row
                if c.value is not None and str(c.value).strip()
            ]
            if cell_values:
                rows.append(cell_values)
        if not rows:
            continue
        max_col = max(len(r) for r in rows)
        md_lines = [f"[시트: {name}]"]
        header = rows[0] + [""] * (max_col - len(rows[0]))
        md_lines.append("| " + " | ".join(header) + " |")
        md_lines.append("| " + " | ".join(["---"] * max_col) + " |")
        for row_data in rows[1:]:
            row_data = row_data + [""] * (max_col - len(row_data))
            md_lines.append("| " + " | ".join(row_data) + " |")
        pages.append("\n".join(md_lines))
    return pages


def build_vectordb(
    parse_fn: Callable,
    chunk_fn: Callable,
    collection_name: str = "eval_documents",
) -> chromadb.Collection:
    """data/docs 하위 PDF·DOCX·XLSX를 파싱해 ChromaDB 컬렉션에 저장한다.

    - parse_fn: Path → list[str] (페이지별 텍스트). PDF에만 적용된다.
    - chunk_fn: (pages, source) → list[dict{id, text, source, page}]
      DOCX는 단일 페이지, XLSX는 시트당 1페이지로 내부 파서가 구성한 뒤
      동일 chunk_fn에 넘겨 전략(A~D)의 청킹 규칙을 그대로 태운다.
    """
    client = chromadb.PersistentClient(
        path=str(CHROMA_DIR),
        settings=Settings(anonymized_telemetry=False),
    )

    try:
        client.delete_collection(collection_name)
    except (ValueError, Exception):
        pass

    collection = client.create_collection(
        name=collection_name,
        metadata={"hnsw:space": "cosine"},
    )

    docs_dir = DATA_DIR / "docs"
    if not docs_dir.exists():
        return collection

    ids: list[str] = []
    documents: list[str] = []
    metadatas: list[dict] = []

    for doc_path in sorted(docs_dir.rglob("*")):
        if not doc_path.is_file():
            continue
        suffix = doc_path.suffix.lower()
        if suffix == ".pdf":
            pages = parse_fn(doc_path)
            fmt = "pdf"
        elif suffix == ".docx":
            pages = _parse_docx_as_pages(doc_path)
            fmt = "docx"
        elif suffix == ".xlsx":
            pages = _parse_xlsx_as_pages(doc_path)
            fmt = "xlsx"
        else:
            continue

        chunks = chunk_fn(pages, doc_path.stem)
        for chunk in chunks:
            ids.append(chunk["id"])
            documents.append(chunk["text"])
            metadatas.append({
                "source": chunk["source"],
                "page": str(chunk["page"]),
                "format": fmt,
            })

    if ids:
        collection.add(ids=ids, documents=documents, metadatas=metadatas)

    return collection


def generate_answer(query: str, context_docs: list[str]) -> str:
    """검색된 컨텍스트를 기반으로 LLM 답변을 생성한다."""
    try:
        import httpx

        context = "\n\n---\n\n".join(context_docs[:3])
        prompt = (
            f"다음 문서 내용을 참고하여 질문에 답하세요.\n\n"
            f"[문서]\n{context}\n\n"
            f"[질문]\n{query}\n\n"
            f"[답변]"
        )

        resp = httpx.post(
            f"{OLLAMA_BASE_URL}/api/generate",
            json={
                "model": OLLAMA_MODEL,
                "prompt": prompt,
                "stream": False,
            },
            timeout=60.0,
        )
        resp.raise_for_status()
        return resp.json().get("response", "")
    except Exception as e:
        return f"[답변 생성 실패: {str(e)[:80]}]"


def load_test_questions() -> list[dict]:
    """test_questions.json을 로드한다."""
    questions_path = DATA_DIR / "test_questions.json"
    if not questions_path.exists():
        return []

    with open(questions_path, encoding="utf-8") as f:
        data = json.load(f)

    return data.get("questions", [])
