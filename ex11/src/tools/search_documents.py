"""ex11 문서 검색 도구 (자식·부모 분리 인덱싱).

- 부모(parent): 페이지 단위 원문. `ex11/data/parent_docs.json`에 캐시
- 자식(child): 200자 내외 청크. ChromaDB에 `parent_id` 메타데이터와 함께 저장

검색 경로는 `src/pipeline.py`의 `RagPipeline`이 조립한다. 이 모듈은
인덱싱·저장소 관리와 기본 툴 진입점(`search_documents`)만 담당한다.
"""

import json
import logging
import os
import uuid
from pathlib import Path

from langchain_core.tools import tool
from rich.console import Console

logger = logging.getLogger(__name__)
console = Console()

BASE_DIR = Path(__file__).resolve().parent.parent.parent  # ex11/
DATA_DIR = BASE_DIR / "data"
DOCS_DIR = DATA_DIR / "docs"
PARENT_DOCS_CACHE = DATA_DIR / "parent_docs.json"

CHILD_CHUNK_SIZE = 220
CHILD_CHUNK_OVERLAP = 40


# ---------------------------------------------------------------------------
# 인덱싱 — 부모 dict + 자식 ChromaDB 두 저장소를 동시에 구축
# ---------------------------------------------------------------------------

def _parse_pdf(file_path: Path, parents: dict[str, dict]) -> None:
    from ..tuning.hybrid_parser import parse_pdf_hybrid

    source = file_path.stem
    strategies: list[str] = []
    with console.status(f"[cyan]PDF 파싱[/] [dim]{file_path.name}[/]", spinner="dots"):
        try:
            for text, page_num, strategy in parse_pdf_hybrid(file_path):
                strategies.append(strategy)
                parent_id = f"{source}__p{page_num}"
                parents[parent_id] = {
                    "content": text,
                    "source": source,
                    "page": page_num,
                    "format": "pdf",
                }
        except Exception as exc:
            console.log(f"[red]✗[/] PDF 파싱 실패: {file_path.name} ({exc})")
            return
    strat_txt = ", ".join(f"p{i+1}={s}" for i, s in enumerate(strategies)) or "(empty)"
    console.log(f"[green]✓[/] {file_path.name} [dim]({strat_txt})[/]")


def _parse_docx(file_path: Path, parents: dict[str, dict]) -> None:
    from docx import Document as DocxDocument

    source = file_path.stem
    with console.status(f"[cyan]DOCX 파싱[/] [dim]{file_path.name}[/]", spinner="dots"):
        try:
            doc = DocxDocument(str(file_path))
            text_parts: list[str] = []
            for para in doc.paragraphs:
                t = para.text.strip()
                if not t:
                    continue
                style = para.style.name
                if style == "Title":
                    text_parts.append(f"# {t}")
                elif style.startswith("Heading"):
                    lvl_str = style.replace("Heading", "").strip()
                    try:
                        lvl = int(lvl_str)
                    except ValueError:
                        lvl = 2
                    text_parts.append(f"{'#' * lvl} {t}")
                elif style == "List Bullet":
                    text_parts.append(f"- {t}")
                else:
                    text_parts.append(t)
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    cells = [c.text.strip() for c in row.cells]
                    text_parts.append("| " + " | ".join(cells) + " |")
                    if i == 0:
                        text_parts.append("| " + " | ".join(["---"] * len(cells)) + " |")
            full = "\n".join(text_parts)
            if full:
                parents[f"{source}__p1"] = {
                    "content": full,
                    "source": source,
                    "page": 1,
                    "format": "docx",
                }
        except Exception as exc:
            console.log(f"[red]✗[/] DOCX 파싱 실패: {file_path.name} ({exc})")
            return
    console.log(f"[green]✓[/] {file_path.name}")


def _parse_xlsx(file_path: Path, parents: dict[str, dict]) -> None:
    import openpyxl

    source = file_path.stem
    sheet_count = 0
    with console.status(f"[cyan]XLSX 파싱[/] [dim]{file_path.name}[/]", spinner="dots"):
        try:
            wb = openpyxl.load_workbook(str(file_path), data_only=True)
            for idx, name in enumerate(wb.sheetnames, start=1):
                ws = wb[name]
                rows: list[list[str]] = []
                for row in ws.iter_rows():
                    cell_values = [
                        str(c.value).strip()
                        for c in row
                        if c.value is not None and str(c.value).strip()
                    ]
                    if cell_values:
                        rows.append(cell_values)
                if not rows:
                    continue
                max_col = max(len(r) for r in rows)
                md_lines = [f"[시트: {name}]"]
                header = rows[0] + [""] * (max_col - len(rows[0]))
                md_lines.append("| " + " | ".join(header) + " |")
                md_lines.append("| " + " | ".join(["---"] * max_col) + " |")
                for row_data in rows[1:]:
                    row_data = row_data + [""] * (max_col - len(row_data))
                    md_lines.append("| " + " | ".join(row_data) + " |")
                parents[f"{source}__s{idx}"] = {
                    "content": "\n".join(md_lines),
                    "source": source,
                    "page": idx,
                    "format": "xlsx",
                }
                sheet_count += 1
        except Exception as exc:
            console.log(f"[red]✗[/] XLSX 파싱 실패: {file_path.name} ({exc})")
            return
    console.log(f"[green]✓[/] {file_path.name} [dim](sheets={sheet_count})[/]")


def _parse_parents() -> dict[str, dict]:
    """data/docs/의 원본을 페이지 단위 부모 문서로 파싱한다.

    반환 형식: {parent_id: {"content": 페이지 전체, "source": str, "page": int, "format": str}}
    PDF는 챕터 10 하이브리드 파서(pypdf → Vision 폴백)를 거친다.
    """
    if not DOCS_DIR.exists():
        return {}

    files = sorted([p for p in DOCS_DIR.rglob("*") if p.is_file()])
    parents: dict[str, dict] = {}
    for file_path in files:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            _parse_pdf(file_path, parents)
        elif suffix == ".docx":
            _parse_docx(file_path, parents)
        elif suffix == ".xlsx":
            _parse_xlsx(file_path, parents)
    return parents


def _split_children(parents: dict[str, dict]) -> list[dict]:
    """부모 문서를 자식 청크로 쪼갠다. 자식은 ChromaDB에 들어갈 단위."""
    children: list[dict] = []
    step = CHILD_CHUNK_SIZE - CHILD_CHUNK_OVERLAP
    for parent_id, parent in parents.items():
        text = parent["content"]
        if not text:
            continue
        start = 0
        while start < len(text):
            chunk = text[start:start + CHILD_CHUNK_SIZE].strip()
            if chunk:
                children.append({
                    "content": chunk,
                    "parent_id": parent_id,
                    "source": parent["source"],
                    "page": str(parent["page"]),
                })
            start += step
    return children


def _save_parent_cache(parents: dict[str, dict]) -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    with open(PARENT_DOCS_CACHE, "w", encoding="utf-8") as f:
        json.dump(parents, f, ensure_ascii=False, indent=2)


def _load_parent_cache() -> dict[str, dict]:
    if not PARENT_DOCS_CACHE.exists():
        return {}
    try:
        with open(PARENT_DOCS_CACHE, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def _build_store() -> tuple[object | None, dict[str, dict]]:
    """ChromaDB 자식 컬렉션 + 부모 dict를 준비해 돌려준다.

    ChromaDB·parent_docs.json이 모두 존재하면 재사용, 없으면 data/docs/에서
    새로 인덱싱한다. 인덱싱 시에는 챕터 10 하이브리드 파서를 거친다.
    """
    try:
        import chromadb
        from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction
    except ImportError:
        console.log("[red]✗[/] chromadb 또는 sentence-transformers 미설치.")
        return None, {}

    chroma_dir = os.getenv("CHROMA_PERSIST_DIR", str(DATA_DIR / "chroma_db"))
    collection_name = os.getenv("CHROMA_COLLECTION_NAME", "connecthr_children")

    with console.status("[cyan]한국어 임베딩 모델 로드 중[/] [dim]ko-sroberta-multitask[/]", spinner="dots"):
        ef = SentenceTransformerEmbeddingFunction(model_name="jhgan/ko-sroberta-multitask")
    console.log("[green]✓[/] 임베딩 모델 준비 완료")

    client = chromadb.PersistentClient(path=chroma_dir)
    parents = _load_parent_cache()

    # 캐시 재사용 조건: 컬렉션 존재 + 첫 레코드에 parent_id + 부모 캐시 존재
    try:
        collection = client.get_collection(collection_name, embedding_function=ef)
        sample = collection.get(limit=1, include=["metadatas"]) if collection.count() > 0 else None
        has_parent_id = bool(
            sample and sample.get("metadatas") and sample["metadatas"][0].get("parent_id")
        )
        if has_parent_id and parents:
            console.log(
                f"[green]✓[/] 기존 인덱싱 재사용 — children={collection.count()}, parents={len(parents)}"
            )
            return collection, parents
        client.delete_collection(collection_name)
    except Exception:
        pass

    # 새 인덱싱
    console.log("[bold cyan]data/docs/에서 자식·부모 인덱싱을 새로 시작합니다[/]")

    parents = _parse_parents()
    if not parents:
        console.log("[yellow]⚠[/] data/docs/가 비었습니다.")
        return None, {}

    with console.status(f"[cyan]자식 청크로 쪼개는 중[/] [dim]{len(parents)}개 부모[/]", spinner="dots"):
        children = _split_children(parents)
    if not children:
        console.log("[red]✗[/] 자식 청크 생성 실패.")
        return None, {}
    console.log(f"[green]✓[/] 자식 청크 {len(children)}개 생성 [dim](부모 {len(parents)}개)[/]")

    with console.status(f"[cyan]ChromaDB에 자식 청크 적재 + 임베딩 계산 중[/]", spinner="dots"):
        collection = client.create_collection(
            name=collection_name,
            embedding_function=ef,
            metadata={"hnsw:space": "cosine"},
        )
        ids = [f"c_{uuid.uuid4().hex[:10]}" for _ in children]
        docs = [c["content"] for c in children]
        metas = [
            {
                "parent_id": c["parent_id"],
                "source": c["source"],
                "page": c["page"],
            }
            for c in children
        ]
        collection.add(ids=ids, documents=docs, metadatas=metas)
        _save_parent_cache(parents)
    console.log(f"[green]✓[/] ChromaDB 적재 완료 [dim]→ {chroma_dir}[/]")

    return collection, parents


# ---------------------------------------------------------------------------
# 싱글톤 접근자
# ---------------------------------------------------------------------------

_STORE_CACHE: tuple[object | None, dict[str, dict]] | None = None


def _get_store() -> tuple[object | None, dict[str, dict]]:
    """(child_collection, parent_docs) 튜플을 1회만 초기화해 재사용."""
    global _STORE_CACHE
    if _STORE_CACHE is None:
        _STORE_CACHE = _build_store()
    return _STORE_CACHE


def _get_vectorstore():
    """자식 청크가 담긴 ChromaDB Collection만 돌려주는 레거시 접근자."""
    collection, _parents = _get_store()
    return collection


def _get_parent_docs() -> dict[str, dict]:
    """부모 문서 dict를 돌려준다."""
    _collection, parents = _get_store()
    return parents


# ---------------------------------------------------------------------------
# 기본 툴 진입점 — 에이전트가 간이 검색용으로 사용
# ---------------------------------------------------------------------------

@tool
def search_documents(query: str):
    """사내 규정·가이드라인·정책 등 비정형 문서 내용을 검색합니다.

    자식 청크에서 벡터 유사도로 상위 3개를 고른 뒤 부모 페이지를 복원해
    돌려줍니다. (더 정밀한 경로는 `src.pipeline.RagPipeline`을 사용하세요.)

    Args:
        query: 검색할 질문 또는 키워드

    Returns:
        부모 문서 단위로 중복 제거된 딕셔너리 목록.
    """
    logger.info("[search_documents] 검색 쿼리: %s", query)

    from ..tuning.parent_doc import ParentDocumentRetriever

    collection, parents = _get_store()
    if collection is None or not parents:
        return []
    retriever = ParentDocumentRetriever(child_collection=collection, parent_docs=parents)
    results = retriever.search(query, top_k=3)
    return [
        {
            "content": r["content"],
            "source": r["source"],
            "score": round(r["score"], 4),
        }
        for r in results
    ]
